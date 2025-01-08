import requests, uuid, json
from docx import Document
import os
from fpdf import FPDF

class Services:

    def translate_transcript(self,text):

        # Add your key and endpoint
        key = ""
        endpoint = "https://api.cognitive.microsofttranslator.com/"

        # location, also known as region.
        # required if you're using a multi-service or regional (not global) resource. It can be found in the Azure portal on the Keys and Endpoint page.
        location = "eastus"

        path = '/translate'
        constructed_url = endpoint + path
        languages = {
            'hi': 'Hindi',
            'mr': 'Marathi',
            'gu': 'Gujarati',
            'bn': 'Bengali',
            'te': 'Telugu',
            'ta': 'Tamil',
            'ur': 'Urdu',
            'kn': 'Kannada',
            'ml': 'Malayalam',
            'pa': 'Punjabi',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'ja': 'Japanese',
            'zh-cn': 'Chinese (Simplified)',
            'ar': 'Arabic'
        }

        params = {
            'api-version': '3.0',
            'from': 'en',
            'to': list(languages.keys())
        }

        headers = {
            'Ocp-Apim-Subscription-Key': key,
            # location required if you're using a multi-service or regional (not global) resource.
            'Ocp-Apim-Subscription-Region': location,
            'Content-type': 'application/json',
            'X-ClientTraceId': str(uuid.uuid4())
        }

        # You can pass more than one object in body.
        body = [{
            'text': text
        }]

        request = requests.post(constructed_url, params=params, headers=headers, json=body)
        response = request.json()


        formatted_response=json.dumps(response, sort_keys=True, ensure_ascii=False, indent=4, separators=(',', ': '))
        # doc = Document()

        # # Add the JSON response to the document
        # doc.add_paragraph("API Response:")
        # doc.add_paragraph(formatted_response)
        # doc.save("response.docx")

        # print("Document created successfully!")
        formatted_response=json.loads(formatted_response)
        return formatted_response

    def multilingual(self,formatted_response):
  
        translations = {}
        '''
        for lang_code in languages:
            print(f"Translating into {languages[lang_code]}...")
            translation = translator.translate(transcript, dest=lang_code).text
            translations[lang_code] = translation
        return translations
        '''
        for item in formatted_response:
                for translation in item["translations"]:
                    translations[translation["to"]] = translation["text"]
            

        return translations
    
    def save_to_pdf(self, file_path, content):
        try:
            # Create directory path if it doesn't exist
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Create PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Configure page settings
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Add content
            pdf.multi_cell(0, 10, content)
            
            # Save the PDF
            pdf.output(file_path)
            return True
            
        except Exception as e:
            print(f"Error saving PDF: {str(e)}")
            return False


if __name__ == "__main__":
    A=Services()
    text="what is a car"
    
    B=A.translate_transcript(text)
    C=A.multilingual(B)
    output_dir = os.path.join(os.getcwd(), "output")
    base_filename = "translated"
    translation_file = os.path.join(output_dir, f"{base_filename}.pdf")
    text=''.join(C.values())

    A.save_to_pdf(translation_file, text) 
    print(C)
    

